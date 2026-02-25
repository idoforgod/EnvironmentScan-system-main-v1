#!/usr/bin/env python3
"""
Generate Word (.docx) report from the special scan markdown file.
"자산 토큰화와 이더리움 가격 폭락" special scanning report.
"""

import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_heading_style(paragraph, level, text):
    """Set heading with Korean-friendly formatting."""
    paragraph.clear()
    run = paragraph.add_run(text)

    if level == 1:
        run.font.size = Pt(18)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
    elif level == 2:
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
    elif level == 3:
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x2C, 0x5F, 0x8E)

    paragraph.paragraph_format.space_before = Pt(12)
    paragraph.paragraph_format.space_after = Pt(6)
    return paragraph

def add_formatted_paragraph(doc, text, bold=False, italic=False, size=11, color=None):
    """Add a formatted paragraph."""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    para.paragraph_format.space_after = Pt(6)
    para.paragraph_format.line_spacing = Pt(18)
    return para

def create_report():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.5)

    # ============================================================
    # COVER PAGE
    # ============================================================

    # Title
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run("붕괴와 구축")
    title_run.font.size = Pt(28)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)

    # Subtitle
    sub_para = doc.add_paragraph()
    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_para.add_run("2026년 2월의 암호자산 대격변")
    sub_run.font.size = Pt(18)
    sub_run.font.bold = False
    sub_run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    doc.add_paragraph()

    # Report info box
    info_para = doc.add_paragraph()
    info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_run = info_para.add_run(
        "이더리움은 28% 폭락했고, 실물자산 토큰화는 315% 성장했다.\n"
        "두 사실이 동시에 진실이다."
    )
    info_run.font.size = Pt(13)
    info_run.font.italic = True
    info_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph()

    # Meta information table
    table = doc.add_table(rows=6, cols=2)
    table.style = 'Light Shading Accent 1'

    meta_data = [
        ("보고서 유형", "특별 스캐닝 — 장편 저널리즘 분석"),
        ("주제", "자산 토큰화와 이더리움 가격 폭락"),
        ("발행일", "2026년 2월 19일"),
        ("데이터 기반", "WF1 Phase 2 완료 — 13개 시그널 (pSST 7.8–9.5)"),
        ("스캔 윈도우", "2026-02-19 기준 24시간 (T₀: 2026-02-20 03:05 UTC)"),
        ("소스", "Bloomberg, CoinGlass, Glassnode, RWA.xyz, DeFiLlama, SEC, Federal Reserve"),
    ]

    for i, (label, value) in enumerate(meta_data):
        row = table.rows[i]
        row.cells[0].text = label
        row.cells[0].paragraphs[0].runs[0].font.bold = True
        row.cells[1].text = value

    doc.add_page_break()

    # ============================================================
    # TABLE OF CONTENTS PLACEHOLDER
    # ============================================================

    toc_heading = doc.add_heading('목 차', level=1)

    toc_items = [
        "도입부: 하나의 날, 두 개의 현실",
        "제1부: 폭락의 해부 — 왜 이더리움은 하루에 28%를 잃었나",
        "  1.1 방아쇠 — 케빈 워시 지명",
        "  1.2 연쇄 청산 — DeFi의 피드백 루프",
        "  1.3 기관 레버리지 — 새로운 증폭자",
        "  1.4 결과 — $2조의 증발",
        "제2부: 구축의 해부 — 왜 기관들은 가격 폭락을 무시하고 인프라를 쌓는가",
        "  2.1 실물자산 토큰화 — $170억의 현실",
        "  2.2 블랙록 BUIDL — 기관 DeFi의 탄생",
        "  2.3 SEC 규제 명확화 — 제도화의 가속",
        "제3부: 역설의 해석 — 두 진실이 공존하는 이유",
        "  3.1 두 개의 DeFi",
        "  3.2 시간 스케일의 불일치",
        "  3.3 ETH 가격 역설",
        "제4부: 시스템 리스크의 새로운 지형",
        "  4.1 TradFi-DeFi 콘테이전 브리지",
        "  4.2 DeFi 청산 메커니즘의 구조적 결함",
        "  4.3 ETH 스테이킹 집중 리스크",
        "제5부: 규제 지형 — 두 갈래 시장의 제도적 기반",
        "  5.1 미국의 규제 명확화 — 이중 트랙",
        "  5.2 이중 시장 구조의 결정화",
        "  5.3 글로벌 규제 다양성",
        "제6부: 전망 — 세 개의 가능한 미래",
        "결론: 두 개의 자산 클래스가 하나의 이름을 공유한다",
        "핵심 모니터링 지표",
    ]

    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(2)
        if item.startswith("  "):
            p.paragraph_format.left_indent = Cm(1)
            p.runs[0].font.size = Pt(10)
        else:
            p.runs[0].font.size = Pt(11)
            p.runs[0].font.bold = True

    doc.add_page_break()

    # ============================================================
    # SIGNAL SUMMARY TABLE
    # ============================================================

    doc.add_heading('WF1 Phase 2 완료 시그널 요약 (13개)', level=1)

    signals_table = doc.add_table(rows=14, cols=4)
    signals_table.style = 'Medium Grid 1 Accent 1'

    # Header
    headers = ['순위', '시그널 제목', 'STEEPs', 'pSST']
    for i, header in enumerate(headers):
        cell = signals_table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].font.bold = True

    signal_data = [
        ('1', 'ETH -28% 폭락/청산 $25억', 'Economic (E)', '9.5'),
        ('2', '케빈 워시 지명 — 긴축 충격', 'Political (P) + Economic (E)', '9.3'),
        ('3', 'RWA $170억 — 315% 급등', 'Economic (E) + Technological (T)', '9.2'),
        ('4', 'DeFi 담보 피드백 루프', 'Economic (E) + Technological (T)', '9.2'),
        ('5', 'Trend Research $20억 레버리지', 'Economic (E)', '9.0'),
        ('6', 'SEC 토큰화 증권 분류', 'Political (P) + Technological (T)', '9.0'),
        ('7', '시장 $2조 증발', 'Economic (E) + Social (S)', '9.0'),
        ('8', 'RWA $240억 온체인 파이프라인', 'Economic (E) + Technological (T)', '8.8 → 8.5'),
        ('9', 'BlackRock BUIDL 멀티체인', 'Economic (E) + Technological (T)', '8.5 → 8.8'),
        ('10', 'ETH/BTC 다년 저점', 'Economic (E) + Technological (T)', '8.7'),
        ('11', 'ETH/BTC 통화 긴축 지표', 'Economic (E)', '8.3'),
        ('12', 'DeFi 기관-투기 이중 구조', 'Technological (T) + Economic (E)', '8.1'),
        ('13', '글로벌 규제 다양성 가속', 'Political (P) + Technological (T)', '7.8'),
    ]

    for i, (rank, title, steeps, psst) in enumerate(signal_data):
        row = signals_table.rows[i+1]
        row.cells[0].text = rank
        row.cells[1].text = title
        row.cells[2].text = steeps
        row.cells[3].text = psst

    doc.add_paragraph()

    doc.add_page_break()

    # ============================================================
    # MAIN CONTENT
    # ============================================================

    # Introduction
    doc.add_heading('도입부: 하나의 날, 두 개의 현실', level=1)

    intro_text = (
        "2026년 2월 19일, 두 개의 사실이 동시에 존재했다.\n\n"
        "이더리움(ETH)은 하루 만에 28% 폭락했다. 24시간 안에 25억 달러의 레버리지 포지션이 "
        "강제 청산됐다. 암호자산 시장 전체 시가총액은 3조 4,000억 달러에서 1조 4,000억 달러로 "
        "2조 달러가 증발했다. 역사상 단일 사건으로는 최대 규모의 명목 자산 소멸이었다.\n\n"
        "동시에, 실물자산(RWA) 토큰화 시장은 전년 동기 대비 315% 성장해 총 170억 달러를 기록했다. "
        "블랙록(BlackRock)의 BUIDL 펀드는 이더리움, 폴리곤, 아발란체, 앱토스 — 네 개 블록체인에 "
        "동시 배포를 확대했다. 240억 달러의 전통 금융 자산이 온체인 이전을 위한 법적·기술적 절차를 "
        "진행 중이었다. 미국 증권거래위원회(SEC)는 토큰화 증권의 규제 분류를 확정했고, 미국 최대 "
        "자산운용사들은 블록체인 기반 결제 인프라에 자본을 투입하고 있었다.\n\n"
        "폭락과 성장. 붕괴와 구축. 두 사실은 모순처럼 보이지만, 같은 하루에 같은 생태계에서 일어났다.\n\n"
        "이 보고서는 왜 이 두 사실이 동시에 진실이며, 그것이 무엇을 의미하는지를 추적한다."
    )

    p = doc.add_paragraph(intro_text)
    p.paragraph_format.line_spacing = Pt(20)
    p.paragraph_format.space_after = Pt(12)

    # ============================================================
    # Part 1
    # ============================================================
    doc.add_heading('제1부: 폭락의 해부 — 왜 이더리움은 하루에 28%를 잃었나', level=1)

    doc.add_heading('1.1 방아쇠 — 케빈 워시 지명', level=2)

    p1_1 = (
        "2026년 2월 19일 오전, 트럼프 행정부는 케빈 워시(Kevin Warsh)를 차기 연방준비제도(Fed) "
        "의장으로 지명했다.\n\n"
        "워시는 낯선 이름이 아니다. 2006년부터 2011년까지 Fed 이사를 역임한 그는 2008년 금융위기 "
        "당시 양적완화(QE)를 '금융 모르핀'이라고 비판했다. 인플레이션에 대한 선제적 대응을 지지했고, "
        "대차대조표의 빠른 축소를 주장해왔다. 시장은 그를 안다. 워시의 Fed는 파월의 Fed가 아니다.\n\n"
        "지명 발표 직후, 시장은 즉각 반응했다. 10년 만기 미국 국채 금리가 18베이시스포인트(bp) "
        "급등해 4.89%를 기록했다. 달러 인덱스(DXY)가 1.4% 강세를 보였다. S&P 500 선물은 2.1% "
        "하락했다. 그리고 암호자산 시장은 발표 당일 15% 폭락했다.\n\n"
        "워시 지명은 단순한 인사 발표가 아니었다. 그것은 통화정책의 패러다임 전환 신호였다. 시장이 "
        "지난 2년간 구축해온 '완화적 환경에서 위험 자산 매수'의 논리 전제가 단번에 허물어졌다."
    )
    doc.add_paragraph(p1_1)

    # Signal box
    signal_box = doc.add_paragraph()
    signal_box.alignment = WD_ALIGN_PARAGRAPH.LEFT
    signal_run = signal_box.add_run(
        "■ 핵심 데이터 | 케빈 워시 지명 충격\n"
        "pSST 9.3 | Political (P) + Economic (E)\n"
        "• 10년 국채 금리: +18bp → 4.89%\n"
        "• DXY: +1.4% (달러 강세)\n"
        "• S&P 500 선물: -2.1%\n"
        "• 암호자산 당일 하락: -15%\n"
        "• 2026년 추가 긴축 시장 반영: +50bp"
    )
    signal_run.font.size = Pt(10)
    signal_run.font.bold = False
    signal_run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
    signal_box.paragraph_format.left_indent = Cm(0.5)
    signal_box.paragraph_format.space_before = Pt(6)
    signal_box.paragraph_format.space_after = Pt(12)

    doc.add_heading('1.2 연쇄 청산 — DeFi의 피드백 루프', level=2)

    p1_2 = (
        "워시 지명이 방아쇠였다면, 탈중앙화 금융(DeFi)의 자동화 메커니즘이 총알이었다.\n\n"
        "DeFi 프로토콜은 담보 비율 하락 시 자동으로 포지션을 청산한다. 설계 의도는 개별 프로토콜의 "
        "지급 능력 보호다. 그러나 시장 전체 차원에서는 이 메커니즘이 정반대로 작동한다.\n\n"
        "구조는 이렇다: ETH 가격 하락 → 담보 비율 하락 → 자동 청산봇 실행 → 추가 매도 압력 → "
        "ETH 가격 추가 하락 → 더 많은 청산 → 반복.\n\n"
        "2026년 2월 19일, 이 피드백 루프가 한 시간 안에 두 번 작동했다. 아베(Aave)는 한 시간 동안 "
        "8억 달러의 청산을 처리했다. 메이커다오(MakerDAO)에서 ETH 담보로 발행된 DAI 21억 달러 어치가 "
        "청산 임계값에 근접했다. 컴파운드(Compound)의 ETH 시장 이용률이 98%로 치솟아 실질적으로 "
        "신규 차입이 동결됐다.\n\n"
        "이것은 DeFi 설계의 구조적 문제다. 개별 프로토콜에는 합리적인 리스크 관리이지만, "
        "시장 전체에는 변동성을 증폭시키는 친경기적(procyclical) 메커니즘이다. 1980년대 "
        "저축대부조합(S&L) 위기와 2008년 모기지 시장이 같은 방식으로 붕괴했다."
    )
    doc.add_paragraph(p1_2)

    # DeFi cascade data box
    defi_box = doc.add_paragraph()
    defi_run = defi_box.add_run(
        "■ DeFi 연쇄 청산 데이터 | pSST 9.2\n"
        "• Aave 청산: 시간당 $8억 (피크)\n"
        "• MakerDAO ETH 담보 위험 노출: $21억\n"
        "• Compound 이용률: 98% (사실상 동결)\n"
        "• DeFi 전체 TVL 감소: -$450억\n"
        "• MEV 봇 청산 포획: 하루 $2,400만"
    )
    defi_run.font.size = Pt(10)
    defi_run.font.color.rgb = RGBColor(0x8B, 0x0000, 0x00)
    defi_box.paragraph_format.left_indent = Cm(0.5)
    defi_box.paragraph_format.space_after = Pt(12)

    doc.add_heading('1.3 기관 레버리지 — 새로운 증폭자', level=2)

    p1_3 = (
        "과거 암호자산 시장의 폭락은 주로 소매 투자자의 레버리지가 원인이었다. 2026년은 달랐다.\n\n"
        "트렌드 리서치(Trend Research)는 2월 18일 기준 암호자산에 20억 달러의 레버리지 롱 포지션을 "
        "보유하고 있었다. 레버리지 비율은 3.5배 — 업계 중간값인 2.1배를 크게 상회했다. "
        "이 단일 기관의 강제 매도가 첫 번째 폭락 파동에서 약 4억 달러의 매도 압력을 기여한 것으로 "
        "추산된다.\n\n"
        "이것이 '제도화의 역설'이다. 기관 투자자의 진입이 시장을 안정시킬 것이라는 기대와 달리, "
        "정교한 레버리지 전략을 가진 기관의 진입은 오히려 폭락의 규모와 속도를 증폭시켰다. "
        "골드만삭스, 모건스탠리 등 주요 프라임 브로커가 헤지펀드에 제공하는 암호자산 레버리지는 "
        "이제 암호시장 변동성이 전통 금융 포트폴리오로 전이되는 새로운 경로가 됐다."
    )
    doc.add_paragraph(p1_3)

    doc.add_heading('1.4 결과 — $2조의 증발', level=2)

    p1_4 = (
        "최종 집계: 암호자산 시장 전체 시가총액이 3조 4,000억 달러에서 1조 4,000억 달러로 감소했다. "
        "2조 달러가 48시간 안에 사라졌다.\n\n"
        "개별 자산 성과: 비트코인 -22%(68,000달러), 이더리움 -28%(2,450달러), "
        "알트코인 인덱스 -45%. ETH/BTC 비율은 0.024로 2020년 초 이후 최저치를 기록했다 — "
        "이더리움이 비트코인 대비 구조적 약세를 시사하는 다년 저점이었다. "
        "비트코인 도미넌스는 58%로 치솟았다."
    )
    doc.add_paragraph(p1_4)

    # Market data table
    market_table = doc.add_table(rows=6, cols=3)
    market_table.style = 'Light Grid Accent 1'

    market_headers = ['지표', '폭락 전', '폭락 후']
    for i, h in enumerate(market_headers):
        market_table.rows[0].cells[i].text = h
        market_table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True

    market_data = [
        ('암호자산 시총', '$3.4조', '$1.4조 (-59%)'),
        ('비트코인 (BTC)', '~$87,000', '$68,000 (-22%)'),
        ('이더리움 (ETH)', '~$3,400', '$2,450 (-28%)'),
        ('ETH/BTC 비율', '0.039', '0.024 (다년 저점)'),
        ('BTC 도미넌스', '52%', '58%'),
    ]

    for i, row_data in enumerate(market_data):
        for j, cell_text in enumerate(row_data):
            market_table.rows[i+1].cells[j].text = cell_text

    doc.add_paragraph()
    doc.add_page_break()

    # ============================================================
    # Part 2
    # ============================================================
    doc.add_heading('제2부: 구축의 해부 — 왜 기관들은 가격 폭락을 무시하고 인프라를 쌓는가', level=1)

    doc.add_heading('2.1 실물자산 토큰화 — $170억의 현실', level=2)

    p2_1 = (
        "같은 날, 실물자산(Real World Asset, RWA) 토큰화 시장의 온체인 총 가치 잠금(TVL)은 "
        "170억 달러를 기록했다. 전년 동기 대비 315% 성장한 수치다.\n\n"
        "RWA 토큰화란 무엇인가? 부동산, 국채, 사모 신용, 원자재 등 전통 금융 자산의 소유권을 "
        "블록체인 위에서 토큰으로 표현하는 것이다. 토큰은 24시간, 365일 거래 가능하고, "
        "스마트 컨트랙트를 통해 수익 분배가 자동화된다. T+2일이 걸리던 결제가 몇 분 안에 처리된다.\n\n"
        "이 숫자들은 선언이 아니다. 법적 구조, 스마트 컨트랙트, 수탁 계약이 이미 체결된 "
        "'실행 중인 자본'이다."
    )
    doc.add_paragraph(p2_1)

    # RWA table
    rwa_table = doc.add_table(rows=6, cols=3)
    rwa_table.style = 'Medium Grid 1 Accent 2'

    rwa_headers = ['자산 카테고리', '온체인 TVL', '파이프라인']
    for i, h in enumerate(rwa_headers):
        rwa_table.rows[0].cells[i].text = h
        rwa_table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True

    rwa_data = [
        ('미국 국채', '$82억', '$82억 이상'),
        ('사모 신용', '$61억', '$61억 이상'),
        ('부동산', '$48억', '$48억 이상'),
        ('원자재', '$49억', '$49억 이상'),
        ('합계', '$170억 (현재)', '$240억 (파이프라인)'),
    ]

    for i, row_data in enumerate(rwa_data):
        for j, cell_text in enumerate(row_data):
            rwa_table.rows[i+1].cells[j].text = cell_text

    rwa_caption = doc.add_paragraph("* 출처: RWA.xyz, McKinsey Digital Assets Report, 2026-02-19")
    rwa_caption.runs[0].font.size = Pt(9)
    rwa_caption.runs[0].font.italic = True

    doc.add_heading('2.2 블랙록 BUIDL — 기관 DeFi의 탄생', level=2)

    p2_2 = (
        "블랙록의 BUIDL(BlackRock USD Institutional Digital Liquidity Fund)은 RWA 토큰화 "
        "시장의 상징이다.\n\n"
        "BUIDL의 현황: 운용자산(AUM) 12억 달러, 투자자 기반 47개 기관, 최소 투자 금액 10만 달러, "
        "T+0 결제(당일 결제), 연간 수익률 5.2%(미국 국채 담보). 그리고 이제 이더리움, 폴리곤, "
        "아발란체, 앱토스 — 네 개 블록체인에 동시 배포.\n\n"
        "멀티체인 확장이 의미하는 것은 단순한 기술 업데이트가 아니다. 블랙록이 단일 블록체인에 "
        "대한 의존도를 낮추고, 체인을 인프라 레이어로 취급하기 시작했다는 선언이다. "
        "어느 체인이 살아남든, 블랙록은 블록체인 기반 금융 상품을 운용할 것이라는 전략적 결정이다.\n\n"
        "BUIDL은 '기관 DeFi'의 최초 완성형 사례다. 전통 금융의 신용 품질(미국 국채)에 블록체인의 "
        "운영 효율성(T+0 결제, 24/7 접근)을 결합했다. 이더리움 가격이 2,450달러든 24,500달러든, "
        "BUIDL의 수익률은 미국 국채 금리에 연동된다."
    )
    doc.add_paragraph(p2_2)

    doc.add_heading('2.3 SEC 규제 명확화 — 제도화의 가속', level=2)

    p2_3 = (
        "2026년 2월 19일, SEC는 토큰화된 전통 증권이 블록체인 표현 방식에 관계없이 기존 증권법의 "
        "적용을 받는다는 예비 가이던스를 발표했다.\n\n"
        "규제 불확실성은 기관 진입의 가장 큰 장벽이었다. SEC 가이던스는 이 질문에 답했다: "
        "토큰은 기존 법 체계 안에 있다. 기관 투자자는 이제 명확한 법적 파라미터 안에서 "
        "토큰화 인프라를 구축할 수 있다.\n\n"
        "그러나 이 규제 명확화는 동시에 진입 장벽을 높인다. 규제 준수 인프라를 갖춘 대형 기관이 "
        "유리하고, 소규모 혁신 프로토콜은 불리하다. 규제 명확화는 기관 RWA 시장을 "
        "과점 구조로 재편할 가능성이 있다."
    )
    doc.add_paragraph(p2_3)

    doc.add_page_break()

    # ============================================================
    # Part 3
    # ============================================================
    doc.add_heading('제3부: 역설의 해석 — 두 진실이 공존하는 이유', level=1)

    doc.add_heading('3.1 두 개의 DeFi', level=2)

    p3_1 = (
        "2026년 2월 19일이 보여준 것은 하나의 생태계 안에 두 개의 평행 현실이 존재한다는 사실이다.\n\n"
        "DeFi 1.0 — 투기적 레이어: 레버리지, 자동화된 청산, 소매 투기, 단기 수익 추구. "
        "이 레이어는 ETH 가격에 고도로 민감하다. 이 레이어가 2월 19일에 붕괴됐다.\n\n"
        "DeFi 2.0 — 기관 인프라 레이어: 규제 준수, 미국 국채 담보, T+0 결제, 기관 수탁, "
        "법적 소유권 구조. 이 레이어는 ETH 가격보다 미국 국채 금리와 규제 명확성에 민감하다. "
        "이 레이어는 2월 19일에도 계속 성장했다."
    )
    doc.add_paragraph(p3_1)

    # Two DeFi comparison table
    defi_comp_table = doc.add_table(rows=6, cols=3)
    defi_comp_table.style = 'Light Grid Accent 1'

    comp_headers = ['특성', 'DeFi 1.0 (투기적)', 'DeFi 2.0 (기관 인프라)']
    for i, h in enumerate(comp_headers):
        defi_comp_table.rows[0].cells[i].text = h
        defi_comp_table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True

    comp_data = [
        ('주요 참여자', '소매 투자자, 헤지펀드', '블랙록, 프랭클린, JP모건'),
        ('담보 자산', 'ETH, 알트코인', '미국 국채, 실물자산'),
        ('핵심 리스크', 'ETH 가격 하락', '규제 변화, 국채 금리'),
        ('시간 스케일', '초~시간 단위', '분기~연 단위'),
        ('2월 19일 결과', '붕괴 (-$2조)', '계속 성장 (+$170억 TVL)'),
    ]

    for i, row_data in enumerate(comp_data):
        for j, cell_text in enumerate(row_data):
            defi_comp_table.rows[i+1].cells[j].text = cell_text

    doc.add_paragraph()

    doc.add_heading('3.2 시간 스케일의 불일치', level=2)

    p3_2 = (
        "두 현실이 공존하는 더 근본적인 이유는 시간 스케일의 불일치다.\n\n"
        "DeFi 1.0은 초(second) 단위로 작동한다. 알고리즘이 자동 청산을 실행하고, 가격이 "
        "분 단위로 움직이고, 투자자들이 시간 안에 포지션을 바꾼다.\n\n"
        "DeFi 2.0은 분기(quarter) 단위로 작동한다. 블랙록이 BUIDL 멀티체인 확장을 결정하는 데 "
        "몇 달이 걸린다. SEC 가이던스가 형성되는 데 몇 년이 걸린다. 240억 달러의 RWA 파이프라인을 "
        "완성하는 데 9개월이 걸린다.\n\n"
        "단기 가격 충격은 DeFi 1.0을 파괴할 수 있다. 그러나 DeFi 2.0의 기관적 결정을 되돌리기에는 "
        "충격이 너무 짧다. 법적 구조, 기술 아키텍처, 규제 관계 — 이것들은 하루짜리 시장 패닉으로 "
        "해체되지 않는다.\n\n"
        "이것이 170억 달러의 RWA TVL이 2조 달러의 시가총액 증발과 동시에 존재할 수 있는 이유다."
    )
    doc.add_paragraph(p3_2)

    doc.add_heading('3.3 ETH 가격 역설', level=2)

    p3_3 = (
        "가장 흥미로운 역설은 이더리움 자체에 있다.\n\n"
        "이더리움은 가장 큰 RWA 토큰화 플랫폼이다. 전체 RWA TVL의 61%가 이더리움 위에 있다. "
        "BUIDL의 핵심 인프라가 이더리움이다. 기관들이 RWA 토큰화를 위해 이더리움 스마트 "
        "컨트랙트를 사용한다.\n\n"
        "그런데 이더리움 토큰(ETH) 가격은 폭락했다.\n\n"
        "이 역설의 해석: 기관들은 이더리움 인프라를 구축하고 있지만, ETH 토큰에 투기하지 않는다. "
        "블랙록은 BUIDL의 가스비로 ETH를 사용하지만, ETH를 장기 보유 자산으로 취급하지 않는다.\n\n"
        "인프라 사용과 토큰 투자는 다르다. 이더리움은 도로망과 같다 — 사람들이 도로를 사용한다고 "
        "해서 도로 부동산에 투자하는 것은 아니다. 단기적으로 ETH 가격은 인프라 유틸리티보다 "
        "거시 통화 조건에 더 민감하게 반응한다."
    )
    doc.add_paragraph(p3_3)

    doc.add_page_break()

    # ============================================================
    # Part 4 - Systemic Risk
    # ============================================================
    doc.add_heading('제4부: 시스템 리스크의 새로운 지형', level=1)

    doc.add_heading('4.1 TradFi-DeFi 콘테이전 브리지', level=2)

    p4_1 = (
        "RWA 토큰화가 진행될수록, 전통 금융(TradFi)과 탈중앙화 금융(DeFi) 사이의 경계가 무너진다. "
        "그 결과는 양방향 전염 경로의 생성이다.\n\n"
        "DeFi → TradFi: DeFi 청산 충격이 블록체인 위의 미국 국채 담보 가치를 훼손한다. "
        "기관 투자자의 토큰화 포트폴리오가 손상된다. 전통 포트폴리오 조정을 위한 마진콜이 발생한다.\n\n"
        "TradFi → DeFi: 워시 지명 같은 전통 금융 이벤트가 DeFi 레버리지 포지션을 청산시킨다. "
        "기관 헤지펀드의 암호자산 강제 매도가 DeFi 프로토콜에 유동성 위기를 일으킨다.\n\n"
        "이것은 금융 역사에서 전례가 없는 구조다. 미국 국채, 상업용 부동산, 사모 신용이 "
        "블록체인 위에 더 많이 올라올수록, 블록체인 특유의 리스크가 전통 자산 시장에 영향을 미친다. "
        "금융 시스템의 새로운 취약점이 탄생하고 있다."
    )
    doc.add_paragraph(p4_1)

    doc.add_heading('4.2 DeFi 청산 메커니즘의 구조적 결함과 해결 경로', level=2)

    p4_2 = (
        "2월 19일의 연쇄 청산은 DeFi의 구조적 결함을 실험실 환경에서 확인했다. "
        "이 결함은 세 가지 해결 경로를 가진다.\n\n"
        "경로 A — 서킷브레이커: 청산 속도나 규모를 일시 중단하는 프로토콜 거버넌스 메커니즘. "
        "시스템 안정성 향상이지만, 탈중앙화 원칙과의 긴장이 있다.\n\n"
        "경로 B — 담보 다양화: ETH 단일 담보에서 RWA(국채, 실물자산)로 다양화. "
        "비상관 자산이 담보가 될수록 피드백 루프가 약화된다. RWA 토큰화가 이 경로의 핵심 재료다.\n\n"
        "경로 C — 기관 리스크 관리 통합: 전통 신용 리스크 프레임워크를 DeFi에 적용. "
        "최대손실액(VaR), 스트레스 테스트, 집중 리스크 한도.\n\n"
        "가장 현실적인 경로는 B와 C의 조합이다. RWA 담보 + 기관 리스크 관리. "
        "이것이 DeFi 2.0의 아키텍처 방향이다."
    )
    doc.add_paragraph(p4_2)

    doc.add_heading('4.3 ETH 스테이킹 집중 리스크', level=2)

    p4_3 = (
        "2월 19일 폭락이 드러낸 또 다른 구조적 위험은 ETH 스테이킹 집중이다.\n\n"
        "리도(Lido)가 전체 스테이킹된 ETH의 32%를 보유한다. 단일 프로토콜이 이더리움 합의 "
        "레이어(consensus layer)의 3분의 1 이상을 통제하는 것은 의도하지 않은 집중화다. "
        "ETH 가격이 계속 하락해 리도에서 대규모 언스테이킹이 발생한다면, 이더리움 네트워크의 "
        "보안 레이어 자체가 일시적으로 약화될 수 있다.\n\n"
        "언스테이킹 대기열 깊이(unstaking queue depth)는 앞으로 60일 동안 가장 중요한 "
        "이더리움 네트워크 건강 지표다."
    )
    doc.add_paragraph(p4_3)

    doc.add_page_break()

    # ============================================================
    # Part 5 - Regulatory Landscape
    # ============================================================
    doc.add_heading('제5부: 규제 지형 — 두 갈래 시장의 제도적 기반', level=1)

    doc.add_heading('5.1 미국의 규제 명확화 — 이중 트랙', level=2)

    p5_1 = (
        "2026년 초 미국의 암호자산 규제는 빠르게 구체화되고 있다. 두 개의 핵심 입법이 진행 중이다.\n\n"
        "GENIUS Act(지니어스법): 스테이블코인 규제 프레임워크. 서클(Circle)의 USDC와 테더(Tether)의 "
        "USDT 같은 달러 연동 스테이블코인에 대한 최소 준비금 요건, 투명성 공개, 인허가 체계를 확립한다. "
        "2026년 7월이 구현 규정 공시 기한이다.\n\n"
        "CLARITY Act(클래리티법): 디지털 자산 시장 구조 법안. 어떤 암호자산이 증권(SEC 관할)이고 "
        "어떤 것이 상품(CFTC 관할)인지를 명확히 한다.\n\n"
        "두 법안의 조합: 기관 투자자가 RWA 토큰화에 대규모로 진입하는 데 필요한 법적 확실성이 "
        "제공된다. SEC 분류 가이던스 + GENIUS Act + CLARITY Act = 제도화의 3각 지지대."
    )
    doc.add_paragraph(p5_1)

    doc.add_heading('5.2 이중 시장 구조의 결정화', level=2)

    p5_2 = (
        "규제 명확화는 이미 진행 중인 시장 분화를 가속화한다.\n\n"
        "Tier 1 — 규제 준수 기관 토큰화: BUIDL, FOBXX, JP모건 오닉스(Onyx). "
        "전통 금융 자산 + 블록체인 효율성. SEC 규제, 기관 수탁, KYC/AML 적용. "
        "이 계층은 수십 조 달러의 전통 자산 이전을 목표로 한다.\n\n"
        "Tier 2 — 허가 없는 DeFi: 유니스왑(Uniswap), 에이브(Aave), 컴파운드(Compound). "
        "소매 사용자, 암호 네이티브 투자자. 가격 변동성과 레버리지. "
        "이 계층은 규제 회색지대에서 작동한다.\n\n"
        "이 이중 구조가 2026-2028년 DeFi 시장의 정의적 특성이 될 것이다."
    )
    doc.add_paragraph(p5_2)

    doc.add_page_break()

    # ============================================================
    # Part 6 - Scenarios
    # ============================================================
    doc.add_heading('제6부: 전망 — 세 개의 가능한 미래', level=1)

    # Scenario A
    scenario_a = doc.add_paragraph()
    run_a = scenario_a.add_run("시나리오 A — '기관 바닥' (확률 45%)")
    run_a.font.bold = True
    run_a.font.size = Pt(12)
    run_a.font.color.rgb = RGBColor(0x00, 0x80, 0x00)

    p_a = (
        "ETH 가격이 2,200-2,400달러에서 바닥을 찾는다. 기관 매수(ETF 축적, RWA 플랫폼의 "
        "ETH 가스비 수요)가 지지선을 형성한다. RWA TVL은 Q2 2026까지 180억 달러로 회복된다. "
        "워시는 확인되지만 시장이 우려했던 것보다 점진적인 대차대조표 축소 타임라인을 제시한다. "
        "암호자산 시장 시가총액이 60일 내 2조 2,000억 달러로 회복된다.\n\n"
        "전제조건: ETH 2,200달러 지지선 유지, 기관 ETF 유입 전환, 워시의 상원 청문회에서 온건한 포워드 가이던스."
    )
    doc.add_paragraph(p_a)

    # Scenario B
    scenario_b = doc.add_paragraph()
    run_b = scenario_b.add_run("시나리오 B — '거시 겨울' (확률 35%)")
    run_b.font.bold = True
    run_b.font.size = Pt(12)
    run_b.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)

    p_b = (
        "워시 확인과 공격적 포워드 가이던스가 지속적인 위험 회피 환경을 만든다. "
        "ETH는 2,200달러 지지를 실패하고 2022년과 유사한 약세장 구조에 진입한다. "
        "RWA 토큰화 파이프라인이 중단되고 DeFi TVL이 850억 달러에서 400억 달러로 하락한다. "
        "BTC 도미넌스가 65%+로 상승하지만, ETH 생태계는 구조적 약세장에 진입한다.\n\n"
        "전제조건: ETH 2,200달러 지지 실패, 워시의 공격적 대차대조표 축소 발표, 기관 RWA 일시 중단 발표."
    )
    doc.add_paragraph(p_b)

    # Scenario C
    scenario_c = doc.add_paragraph()
    run_c = scenario_c.add_run("시나리오 C — '분화 회복' (확률 20%)")
    run_c.font.bold = True
    run_c.font.size = Pt(12)
    run_c.font.color.rgb = RGBColor(0xFF, 0x88, 0x00)

    p_c = (
        "가장 복잡한 결과: 기관 DeFi(BUIDL, 규제 준수 RWA)는 연말 300억+에 도달하는 반면, "
        "투기적 DeFi는 억압된 상태를 유지한다. ETH 가격은 적당히 회복(2,800-3,200달러 범위)하고, "
        "RWA 유틸리티가 바닥을 제공한다. 소매 암호자산 채택은 억압되어 있지만, "
        "기관 채택이 전체 생태계의 구조적으로 더 높은 바닥을 만든다."
    )
    doc.add_paragraph(p_c)

    doc.add_page_break()

    # ============================================================
    # Conclusion
    # ============================================================
    doc.add_heading('결론: 두 개의 자산 클래스가 하나의 이름을 공유한다', level=1)

    conclusion_text = (
        "2026년 2월 19일이 보여준 것은 '암호자산'이라는 단어 아래 두 개의 근본적으로 다른 "
        "자산 클래스가 존재한다는 것이다.\n\n"
        "하나는 투기적 암호자산이다. 레버리지, 자동화된 청산, 거시 통화 조건에 극도로 민감한 자산. "
        "이것은 2014년부터 존재했고, 2022년 붕괴와 2026년 붕괴를 반복하고 있다.\n\n"
        "다른 하나는 블록체인 기반 금융 인프라다. 미국 국채 토큰, 규제 준수 결제 시스템, "
        "기관 수탁, T+0 결제. 이것은 2024-2026년에 걸쳐 처음으로 기관 질량을 획득했다.\n\n"
        "두 자산 클래스는 같은 블록체인 네트워크를 사용하고, 같은 '이더리움' 또는 'DeFi'라는 "
        "이름을 공유한다. 그러나 그것들은 다른 물리학으로, 다른 시간 스케일로, 다른 투자자 "
        "기반으로 작동한다.\n\n"
        "단기 가격이 장기 인프라 구축을 평가하는 데 잘못된 렌즈라는 것 — "
        "이것이 2026년 2월 19일의 가장 중요한 신호다.\n\n"
        "ETH 가격이 2,450달러에서 회복될지 더 하락할지는 중요하다. 그러나 더 중요한 질문은: "
        "2028년, 2030년에 전통 금융 자산의 어느 정도 비율이 블록체인 위에 존재할 것인가? "
        "그리고 그 결정의 수혜자는 ETH 토큰 보유자인가, 아니면 블록체인 인프라 위에 서비스를 "
        "제공하는 전통 금융 기관인가?\n\n"
        "이 질문에 대한 답이 향후 5년 암호자산 시장의 구조를 결정한다."
    )
    doc.add_paragraph(conclusion_text)

    # ============================================================
    # Monitoring Indicators
    # ============================================================
    doc.add_heading('핵심 모니터링 지표', level=1)

    mon_table = doc.add_table(rows=10, cols=3)
    mon_table.style = 'Light Grid Accent 1'

    mon_headers = ['시간 범위', '지표', '임계값/해석']
    for i, h in enumerate(mon_headers):
        mon_table.rows[0].cells[i].text = h
        mon_table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True

    mon_data = [
        ('48시간', 'ETH 가격 vs $2,200 지지선', '하향 돌파 = 추가 청산 파동 위험'),
        ('48시간', '이더리움 ETF 일별 자금 흐름', '순유입 전환 = 기관 바닥 신호'),
        ('48시간', '리도(Lido) 언스테이킹 대기열', '급증 = 네트워크 보안 리스크'),
        ('30일', '워시 상원 인사청문회 발언', '대차대조표 축소 타임라인 핵심'),
        ('30일', 'RWA 파이프라인 일시 중단/가속화', '기관 신뢰 지표'),
        ('30일', 'DeFi 거버넌스 청산 개선 제안', 'DeFi 2.0 전환 신호'),
        ('90일', 'RWA TVL 방향성', '$180억 이상 = 기관 수요 확인'),
        ('90일', 'BUIDL AUM 궤적', '주간 성장률 모니터링'),
        ('90일', 'SEC 면제 신청 처리 결과', '기관 토큰화 시장 접근성'),
    ]

    for i, row_data in enumerate(mon_data):
        for j, cell_text in enumerate(row_data):
            mon_table.rows[i+1].cells[j].text = cell_text

    doc.add_paragraph()

    # ============================================================
    # Footer note
    # ============================================================
    footer_para = doc.add_paragraph()
    footer_run = footer_para.add_run(
        "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
        "이 보고서는 Triple Environmental Scanning System v2.5.0 (WF1 특별 스캐닝) 기반으로 생성됐습니다.\n"
        "스캔 일시: 2026-02-19 | 스캔 윈도우: 24시간 | 시그널 수: 13개 | pSST 범위: 7.8–9.5\n"
        "데이터 출처: Bloomberg, CoinGlass, Glassnode, RWA.xyz, DeFiLlama, SEC, Federal Reserve\n"
        "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━"
    )
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Save
    output_path = "/Users/cys/Desktop/AIagentsAutomation/EnvironmentScan-system-main-v4/env-scanning/integrated/reports/daily/특별스캔-자산토큰화-이더리움-2026-02-19.docx"
    doc.save(output_path)
    print(f"Word report saved to: {output_path}")
    return output_path

if __name__ == "__main__":
    path = create_report()
    print(f"SUCCESS: {path}")
